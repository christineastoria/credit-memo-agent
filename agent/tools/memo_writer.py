"""
Memo writer tool for generating branded Word documents.

Takes the assembled credit memo content and produces a formatted .docx
file with proper headings, tables, and a branded header. Uses python-docx
for document generation.
"""

import os
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from langchain.tools import tool

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "..", "output")


def _add_branded_header(doc, borrower_name: str, analyst: str):
    """Add a branded header block to the top of the document."""
    # Title
    title = doc.add_heading("CREDIT INVESTMENT MEMO", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)  # Navy blue

    # Subtitle with borrower name
    subtitle = doc.add_heading(borrower_name, level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata line
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(
        f"Prepared by: {analyst}  |  Date: {datetime.now().strftime('%B %d, %Y')}  |  CONFIDENTIAL"
    )
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(128, 128, 128)

    # Divider line
    doc.add_paragraph("_" * 72)


def _add_section(doc, title: str, content: str):
    """Add a titled section with body text."""
    doc.add_heading(title, level=2)
    doc.add_paragraph(content)


def _add_metrics_table(doc, metrics: dict):
    """Add a formatted table of credit metrics."""
    doc.add_heading("Key Credit Metrics", level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Shading Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = "Metric"
    header_cells[1].text = "Value"

    # Data rows
    for metric_name, metric_value in metrics.items():
        row_cells = table.add_row().cells
        row_cells[0].text = str(metric_name)
        row_cells[1].text = str(metric_value)


def _add_citations(doc, citations: list):
    """Add a citations/sources section at the end."""
    doc.add_heading("Sources & Citations", level=2)
    for i, citation in enumerate(citations, 1):
        doc.add_paragraph(f"[{i}] {citation}", style="List Number")


@tool
def generate_memo_docx(
    borrower_name: str,
    analyst: str,
    executive_summary: str,
    business_overview: str,
    capital_structure: str,
    credit_metrics: dict,
    industry_analysis: str,
    downside_scenarios: str,
    recommendation: str,
    citations: list,
    disclaimer: str,
) -> str:
    """Generate a formatted Word document (.docx) for the credit investment memo.

    Args:
        borrower_name: Name of the borrower/company being analyzed
        analyst: Name of the analyst preparing the memo
        executive_summary: Brief summary of the credit opportunity and recommendation
        business_overview: Description of the company, its operations, and market position
        capital_structure: Current debt structure, maturity profile, and covenants
        credit_metrics: Dict of metric name -> value (e.g., {"Leverage (Debt/EBITDA)": "4.7x"})
        industry_analysis: Sector dynamics, competitive position, and market trends
        downside_scenarios: Stress test results and downside case analysis
        recommendation: Final investment recommendation with rationale
        citations: List of source strings for the citations appendix
        disclaimer: Compliance disclaimer text

    Returns:
        Path to the generated .docx file
    """
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Build the document section by section
    _add_branded_header(doc, borrower_name, analyst)
    _add_section(doc, "Executive Summary", executive_summary)
    _add_section(doc, "Business Overview", business_overview)
    _add_section(doc, "Capital Structure", capital_structure)
    _add_metrics_table(doc, credit_metrics)
    _add_section(doc, "Industry Analysis", industry_analysis)
    _add_section(doc, "Downside Scenarios", downside_scenarios)
    _add_section(doc, "Recommendation", recommendation)
    _add_citations(doc, citations)

    # Add compliance disclaimer at the bottom
    doc.add_paragraph("")
    disclaimer_para = doc.add_paragraph()
    disclaimer_run = disclaimer_para.add_run(disclaimer)
    disclaimer_run.font.size = Pt(8)
    disclaimer_run.font.color.rgb = RGBColor(128, 128, 128)
    disclaimer_run.font.italic = True

    # Save with a timestamped filename
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_name = borrower_name.lower().replace(" ", "_")
    filename = f"credit_memo_{safe_name}_{timestamp}.docx"
    filepath = os.path.join(OUTPUT_DIR, filename)
    doc.save(filepath)

    return f"Memo saved to: {filepath}"
